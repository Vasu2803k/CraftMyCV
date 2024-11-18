from pathlib import Path
import tempfile
import os
import warnings
from typing import Type
from io import StringIO
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams
from docx import Document
import olefile
from crewai.tools import BaseTool
from pydantic import BaseModel, Field
import traceback

# Ignore warnings
warnings.filterwarnings('ignore')

class TextExtractionInput(BaseModel):
    """Input schema for TextExtractionTool."""
    file_path: str = Field(
        ..., 
        description="Path to the file from which to extract text. Supported formats: PDF, DOCX, DOC, TXT, PNG, JPG, JPEG"
    )

class TextExtractor:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.extract_text_from_pdf,
            '.docx': self.extract_text_from_docx,
            '.doc': self.extract_text_from_doc,
            '.txt': self.extract_text_from_txt,
            '.png': self.perform_ocr,
            '.jpg': self.perform_ocr,
            '.jpeg': self.perform_ocr
        }
        
        # Check for required dependencies
        try:
            pytesseract.get_tesseract_version()
        except Exception:
            warnings.warn(
                "Tesseract is not installed or not in PATH. OCR functionality will not work. "
                "Please install Tesseract: https://github.com/UB-Mannheim/tesseract/wiki"
            )

    def perform_text_extraction(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension with fallback to OCR.
        
        Args:
            file_path (str): Path to the file to extract text from
            
        Returns:
            str: Extracted text if successful
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ValueError: If the file type is not supported
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
        try:
            # Try primary extraction method
            extraction_method = self.supported_formats[file_extension]
            text = extraction_method(file_path)
            
            # If primary extraction fails or returns minimal text, try OCR fallback
            if not text or len(text.strip()) < 50:  # Threshold for minimum text length
                if file_extension in ['.pdf', '.docx', '.doc']:
                    print(f"Primary extraction yielded insufficient text. Attempting OCR fallback for {file_path}")
                    text = self.ocr_fallback(file_path, file_extension)
            
            if not text:
                raise ValueError(f"No text extracted from {file_path}")
            return text.strip()
            
        except Exception as e:
            print(f"Stack trace:\n{traceback.format_exc()}")
            # If primary extraction fails, try OCR fallback for supported formats
            if file_extension in ['.pdf', '.docx', '.doc']:
                print(f"Primary extraction failed. Attempting OCR fallback for {file_path}")
                try:
                    text = self.ocr_fallback(file_path, file_extension)
                    if text:
                        return text.strip()
                except Exception as ocr_e:
                    raise ValueError(f"Both primary extraction and OCR fallback failed: {str(e)} | OCR Error: {str(ocr_e)}")
            raise ValueError(f"Error extracting text from {file_path}: {str(e)}")

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from a PDF file using pdfminer.six"""
        path = Path(pdf_path)
        if not path.exists():
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        if not pdf_path.lower().endswith('.pdf'):
            raise ValueError(f"Invalid file format. Expected PDF, got: {path.suffix}")
        
        output_string = StringIO()
        try:
            with open(pdf_path, 'rb') as file:
                laparams = LAParams(
                    line_margin=0.5,
                    word_margin=0.1,
                    char_margin=2.0,
                    boxes_flow=0.5,
                    detect_vertical=True
                )
                extract_text_to_fp(file, output_string, laparams=laparams)
                text = output_string.getvalue()
                return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {e}")
        finally:
            output_string.close()

    def extract_text_from_doc(self, doc_path: str) -> str:
        """Extract text from a DOC file using olefile"""
        path = Path(doc_path)
        if not path.exists():
            raise FileNotFoundError(f"DOC file not found: {doc_path}")
        if not doc_path.lower().endswith('.doc'):
            raise ValueError(f"Invalid file format. Expected DOC, got: {path.suffix}")
        
        ole = None
        try:
            with open(doc_path, 'rb') as file:
                ole = olefile.OleFileIO(file)
                if not ole.exists('WordDocument'):
                    raise ValueError("No 'WordDocument' stream found in the DOC file.")
                
                # Try multiple streams in order of preference
                for stream in ['WordDocument', 'BodyText']:
                    if ole.exists(stream):
                        raw_text = ole.openstream(stream).read()
                        # Try multiple encodings
                        for encoding in ['utf-8', 'latin-1', 'cp1252']:
                            try:
                                text = raw_text.decode(encoding, errors='ignore')
                                if text.strip():
                                    return text.strip()
                            except UnicodeDecodeError:
                                continue
                
                raise ValueError("No readable text found in any DOC file stream")
        except Exception as e:
            raise ValueError(f"Error extracting text from DOC: {e}")
        finally:
            if ole:
                ole.close()

    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from a TXT file"""
        path = Path(txt_path)
        if not path.exists():
            raise FileNotFoundError(f"TXT file not found: {txt_path}")
        if not txt_path.lower().endswith('.txt'):
            raise ValueError(f"Invalid file format. Expected TXT, got: {path.suffix}")
        
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
                if not text.strip():
                    raise ValueError("No text extracted from the TXT file.")
                return text.strip()
        except UnicodeDecodeError:
            # Fallback to different encoding if UTF-8 fails
            with open(txt_path, 'r', encoding='latin-1') as file:
                text = file.read()
                if not text.strip():
                    raise ValueError("No text extracted from the TXT file.")
                return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from TXT: {e}")
        finally:
            if 'file' in locals():
                file.close()
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from a DOCX file using python-docx"""
        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        if not docx_path.lower().endswith('.docx'):
            raise ValueError(f"Invalid file format. Expected DOCX, got: {path.suffix}")
        
        try:
            doc = Document(docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            if not text.strip():
                raise ValueError("No text extracted from the DOCX file.")
            return text.strip()
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {e}")
        
    def perform_ocr(self, image_path: str) -> str:
        """Perform OCR on an image file using pytesseract"""
        path = Path(image_path)
        if not path.exists():
            raise FileNotFoundError(f"Image file not found: {image_path}")
        if not image_path.lower().endswith(('.png', '.jpg', '.jpeg')):
            raise ValueError(f"Invalid file format. Expected image file, got: {path.suffix}")
        
        try:
            return self._optimize_and_ocr_image(image_path)
        except Exception as e:
            raise ValueError(f"Error performing OCR: {e}")

    def ocr_fallback(self, file_path: str, file_extension: str) -> str:
        """Perform OCR on document as fallback method."""
        if not file_extension in ['.pdf', '.docx', '.doc']:
            raise ValueError(f"OCR fallback not supported for {file_extension}")
        
        try:
            if file_extension == '.pdf':
                text = self._ocr_pdf(file_path)
            else:  # .docx or .doc
                text = self._ocr_document(file_path)
            
            if not text or len(text.strip()) < 20:  # Minimum threshold for meaningful text
                raise ValueError("OCR produced insufficient text")
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"OCR fallback failed for {file_extension} file: {str(e)}")

    def _optimize_and_ocr_image(self, image_path: str) -> str:
        """Optimize image and perform OCR with enhanced settings."""
        path = Path(image_path)
        if not path.exists():
            raise FileNotFoundError(f"Image file not found: {image_path}")
        
        try:
            image = Image.open(image_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                try:
                    image = image.convert('RGB')
                except Exception as e:
                    raise ValueError(f"Failed to convert image to RGB mode: {str(e)}")
            
            # Image preprocessing
            try:
                max_dimension = 4500
                if max(image.size) > max_dimension:
                    ratio = max_dimension / max(image.size)
                    new_size = tuple(int(dim * ratio) for dim in image.size)
                    image = image.resize(new_size, Image.Resampling.LANCZOS)
            except Exception as e:
                raise ValueError(f"Image resize failed: {str(e)}")
            
            try:
                # Configure OCR settings for better accuracy
                custom_config = r'--oem 3 --psm 6 --dpi 300'
                text = pytesseract.image_to_string(image, config=custom_config)
                
                if not text.strip():
                    # Try alternative OCR settings if no text was found
                    alternative_config = r'--oem 3 --psm 3 --dpi 300'
                    text = pytesseract.image_to_string(image, config=alternative_config)
                
                return text.strip()
            except Exception as e:
                raise ValueError(f"OCR process failed: {str(e)}")
        except Exception as e:
            raise ValueError(f"Image OCR optimization failed: {str(e)}")
        finally:
            if 'image' in locals():
                image.close()

    def _ocr_pdf(self, pdf_path: str) -> str:
        """Convert PDF to images and perform OCR."""
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                # Process pages in batches to manage memory
                text_parts = []
                batch_size = 10  # Process 10 pages at a time
                
                # Get total number of pages
                images = convert_from_path(pdf_path, first_page=1, last_page=1)
                total_pages = len(convert_from_path(pdf_path, first_page=1, last_page=None))
                
                for start_page in range(1, total_pages + 1, batch_size):
                    end_page = min(start_page + batch_size - 1, total_pages)
                    batch_images = convert_from_path(
                        pdf_path,
                        first_page=start_page,
                        last_page=end_page
                    )
                    
                    for i, image in enumerate(batch_images):
                        temp_image_path = os.path.join(temp_dir, f'page_{start_page + i}.png')
                        image.save(temp_image_path, 'PNG')
                        text = self._optimize_and_ocr_image(temp_image_path)
                        if text:
                            text_parts.append(text)
                        # Clean up the temporary image right away
                        os.remove(temp_image_path)

                return '\n\n'.join(text_parts)
            except Exception as e:
                raise ValueError(f"PDF OCR failed: {str(e)}")

    def _ocr_document(self, doc_path: str) -> str:
        """Convert DOCX to images and perform OCR.
        
        Args:
            doc_path (str): Path to the DOCX file
            
        Returns:
            str: Extracted text from OCR
            
        Raises:
            ValueError: If conversion or OCR fails
        """
        if not doc_path.lower().endswith('.docx'):
            raise ValueError("This method only supports DOCX files")
        
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                # Extract images from DOCX
                doc = Document(doc_path)
                text_parts = []
                
                # First try to get text directly from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # Extract and process any images in the document
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_blob = rel.target_part.blob
                            temp_image_path = os.path.join(temp_dir, f'image_{len(text_parts)}.png')
                            
                            with open(temp_image_path, 'wb') as img_file:
                                img_file.write(image_blob)
                                
                            # Perform OCR on the extracted image
                            image_text = self._optimize_and_ocr_image(temp_image_path)
                            if image_text.strip():
                                text_parts.append(image_text)
                        except Exception as img_e:
                            print(f"Warning: Failed to process embedded image: {str(img_e)}")
                            continue
                
                if not text_parts:
                    raise ValueError("No text or readable images found in the DOCX file")
                    
                return '\n\n'.join(text_parts)
                
            except Exception as e:
                raise ValueError(f"DOCX OCR failed: {str(e)}")

class TextExtractionTool(BaseTool):
    name: str = "Text Extraction Tool"
    description: str = (
        "A tool for extracting text from various file formats including PDF, DOCX, DOC, TXT, "
        "and images (PNG, JPG, JPEG). Features robust error handling and OCR fallback."
    )
    args_schema: Type[BaseModel] = TextExtractionInput
    
    def __init__(self):
        # Initialize the parent class first
        super().__init__()
        # Initialize the extractor as a public attribute
        self._extractor = TextExtractor()

    def _run(self, file_path: str) -> str:
        """Execute the text extraction tool."""
        try:
            text = self._extractor.perform_text_extraction(file_path)
            if not text:
                return "Warning: No text could be extracted from the file."
            return text
        except FileNotFoundError as e:
            print(f"Stack trace:\n{traceback.format_exc()}")
            return f"File not found error: {str(e)}"
        except ValueError as e:
            print(f"Stack trace:\n{traceback.format_exc()}")
            return f"Validation error: {str(e)}"
        except Exception as e:
            print(f"Stack trace:\n{traceback.format_exc()}")
            return f"Unexpected error during text extraction: {str(e)}"